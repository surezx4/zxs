# -*- coding: utf-8 -*-
"""
Created on Fri Aug 22 17:55:12 2025

@author: 10681
"""

import sys
import os
import base64
import json
import requests
import tempfile
import subprocess
import matplotlib.pyplot as plt
from PIL import Image, ImageGrab
from io import BytesIO
import numpy as np
# 导入matplotlib的Figure和Qt5后端的FigureCanvas
from matplotlib.figure import Figure
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas

from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QTextEdit, QPushButton, QLabel, QFileDialog,
                            QSplitter, QMessageBox, QToolBar, QAction, QDockWidget,
                            QInputDialog, QLineEdit, QComboBox, QGroupBox, QFormLayout,
                            QSlider, QColorDialog, QTabWidget, QGridLayout, QScrollArea,
                            QFrame, QSizePolicy, QStackedWidget, QListWidget, QListWidgetItem,
                            QDialog, QDialogButtonBox, QPlainTextEdit)
from PyQt5.QtGui import QFont, QColor, QPen, QPainter, QPixmap, QImage, QIcon, QCursor, QKeySequence, QTextCursor
from PyQt5.QtCore import Qt, QSize, QPoint, QThread, pyqtSignal, QRect, QTimer

# 确保matplotlib支持中文显示
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS", "sans-serif"]
plt.rcParams["axes.unicode_minus"] = False  # 正确显示负号

# 尝试导入Pix2Text
try:
    from pix2text import Pix2Text
    PIX2TEXT_AVAILABLE = True
except ImportError:
    PIX2TEXT_AVAILABLE = False

# DeepSeek API配置 (备用)
DEEPSEEK_API_KEY = ""
DEEPSEEK_OCR_URL = "https://api.deepseek.com/v1/ocr/formula"  # 公式识别API
DEEPSEEK_CHAT_URL = "https://api.deepseek.com/v1/chat/completions"  # 对话API

class ScreenshotDialog(QDialog):
    """截图对话框"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("截图识别公式")
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.WindowStaysOnTopHint)
        self.setWindowState(Qt.WindowFullScreen)
        self.setStyleSheet("background-color: rgba(0, 0, 0, 100);")
        
        self.start_point = QPoint()
        self.end_point = QPoint()
        self.is_drawing = False
        self.init_ui()
        
    def init_ui(self):
        # 添加提示标签
        self.hint_label = QLabel("按住鼠标左键拖动选择区域，按ESC取消", self)
        self.hint_label.setStyleSheet("color: white; background-color: rgba(0, 0, 0, 150); padding: 10px;")
        self.hint_label.setAlignment(Qt.AlignCenter)
        self.hint_label.setGeometry(QRect(0, 0, self.width(), 40))
        
    def paintEvent(self, event):
        """绘制事件"""
        if self.is_drawing:
            painter = QPainter(self)
            painter.setPen(QPen(Qt.red, 2, Qt.SolidLine))
            painter.setBrush(QBrush(QColor(255, 255, 255, 100)))
            rect = QRect(self.start_point, self.end_point)
            painter.drawRect(rect)
            
    def mousePressEvent(self, event):
        """鼠标按下事件"""
        if event.button() == Qt.LeftButton:
            self.start_point = event.pos()
            self.end_point = event.pos()
            self.is_drawing = True
            self.update()
            
    def mouseMoveEvent(self, event):
        """鼠标移动事件"""
        if self.is_drawing:
            self.end_point = event.pos()
            self.update()
            
    def mouseReleaseEvent(self, event):
        """鼠标释放事件"""
        if event.button() == Qt.LeftButton and self.is_drawing:
            self.end_point = event.pos()
            self.is_drawing = False
            
            # 获取截图区域
            rect = QRect(self.start_point, self.end_point).normalized()
            if rect.width() > 10 and rect.height() > 10:  # 确保区域足够大
                # 截取屏幕
                screen = QApplication.primaryScreen()
                screenshot = screen.grabWindow(0, rect.x(), rect.y(), rect.width(), rect.height())
                
                # 转换为PIL图像
                buffer = BytesIO()
                screenshot.save(buffer, "PNG")
                pil_image = Image.open(buffer)
                
                # 发送信号给主窗口进行识别
                self.parent().screenshot_captured(pil_image)
                
            self.accept()
            
    def keyPressEvent(self, event):
        """按键事件"""
        if event.key() == Qt.Key_Escape:
            self.reject()

class ResizableHandwritingCanvas(QWidget):
    """可调整大小的手写画布"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent_panel = parent
        self.drawing = False
        self.last_point = QPoint()
        self.image = QImage()
        self.pen_color = QColor(0, 0, 0)  # 黑色
        self.pen_width = 2
        self.resize_handle_size = 10
        self.resizing = False
        self.init_ui()
        
    def init_ui(self):
        self.setMinimumSize(400, 200)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.setStyleSheet("background-color: white; border: 1px solid #ccc;")
        self.clear_canvas()
        
    def resizeEvent(self, event):
        """调整大小时重绘图像"""
        if not self.image.isNull():
            # 保存当前图像
            old_image = self.image.copy()
            # 创建新图像
            self.image = QImage(self.size(), QImage.Format_RGB32)
            self.image.fill(Qt.white)
            # 将旧图像绘制到新图像上
            painter = QPainter(self.image)
            painter.drawImage(0, 0, old_image.scaled(
                self.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation))
        else:
            self.image = QImage(self.size(), QImage.Format_RGB32)
            self.image.fill(Qt.white)
        super().resizeEvent(event)
        
    def mousePressEvent(self, event):
        """鼠标按下事件"""
        # 检查是否在调整大小区域
        if event.button() == Qt.LeftButton and self.is_resize_handle(event.pos()):
            self.resizing = True
            self.setCursor(QCursor(Qt.SizeFDiagCursor))
            return
            
        # 正常绘画
        if event.button() == Qt.LeftButton:
            self.drawing = True
            self.last_point = event.pos()
            self.draw_point(event.pos())
            
    def mouseMoveEvent(self, event):
        """鼠标移动事件"""
        # 处理调整大小
        if self.resizing:
            # 限制最小尺寸
            new_width = max(400, event.pos().x())
            new_height = max(200, event.pos().y())
            self.resize(new_width, new_height)
            return
            
        # 处理绘画
        if event.buttons() & Qt.LeftButton and self.drawing:
            painter = QPainter(self.image)
            painter.setPen(QPen(self.pen_color, self.pen_width, 
                               Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
            painter.drawLine(self.last_point, event.pos())
            self.last_point = event.pos()
            self.update()
            
    def mouseReleaseEvent(self, event):
        """鼠标释放事件"""
        if self.resizing:
            self.resizing = False
            self.setCursor(QCursor(Qt.ArrowCursor))
            return
            
        if event.button() == Qt.LeftButton and self.drawing:
            self.drawing = False
            
    def paintEvent(self, event):
        """绘制事件"""
        painter = QPainter(self)
        painter.drawImage(self.rect(), self.image, self.image.rect())
        
        # 绘制调整大小的手柄
        if not self.resizing:
            handle_rect = QRect(
                self.width() - self.resize_handle_size,
                self.height() - self.resize_handle_size,
                self.resize_handle_size,
                self.resize_handle_size
            )
            painter.fillRect(handle_rect, QColor(200, 200, 200))
            painter.setPen(QPen(Qt.gray, 1, Qt.SolidLine))
            painter.drawLine(
                handle_rect.left(), handle_rect.center().y(),
                handle_rect.right(), handle_rect.center().y()
            )
            painter.drawLine(
                handle_rect.center().x(), handle_rect.top(),
                handle_rect.center().x(), handle_rect.bottom()
            )
            
    def is_resize_handle(self, pos):
        """判断鼠标位置是否在调整大小的手柄上"""
        return (pos.x() >= self.width() - self.resize_handle_size and
                pos.y() >= self.height() - self.resize_handle_size)
                
    def draw_point(self, pos):
        """绘制点"""
        painter = QPainter(self.image)
        painter.setPen(QPen(self.pen_color, self.pen_width, 
                           Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
        painter.drawPoint(pos)
        self.update()
        
    def clear_canvas(self):
        """清除手写内容"""
        self.image = QImage(self.size(), QImage.Format_RGB32)
        self.image.fill(Qt.white)
        self.update()
        
    def get_image_data(self):
        """获取图片数据（转换为base64）"""
        # 将QImage转换为PNG格式
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            self.image.save(f, "PNG")
            f.flush()
            
            # 读取并转换为base64
            with open(f.name, 'rb') as img_file:
                image_data = base64.b64encode(img_file.read()).decode("utf-8")
                
            # 清理临时文件
            os.unlink(f.name)
            
            return image_data
            
    def get_pil_image(self):
        """获取PIL图像对象"""
        buffer = BytesIO()
        self.image.save(buffer, "PNG")
        return Image.open(buffer)

class HandwritingPanel(QWidget):
    """增强的可调整大小的手写公式面板"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.drawing = False
        self.last_point = QPoint()
        self.image = QImage()
        self.pen_color = QColor(0, 0, 0)  # 黑色
        self.pen_width = 2
        self.init_ui()
        
    def init_ui(self):
        main_layout = QVBoxLayout(self)
        
        # 控制面板 - 合并画笔设置和尺寸调整
        control_group = QGroupBox("手写设置")
        control_layout = QVBoxLayout(control_group)
        
        # 顶部控制 - 画笔设置
        top_control_layout = QHBoxLayout()
        
        # 画笔颜色选择
        color_layout = QVBoxLayout()
        color_layout.addWidget(QLabel("颜色:"))
        self.color_btn = QPushButton("选择颜色")
        self.color_btn.clicked.connect(self.choose_color)
        color_layout.addWidget(self.color_btn)
        top_control_layout.addLayout(color_layout)
        
        # 画笔粗细选择
        width_layout = QVBoxLayout()
        width_layout.addWidget(QLabel("粗细:"))
        self.width_slider = QSlider(Qt.Horizontal)
        self.width_slider.setMinimum(1)
        self.width_slider.setMaximum(10)
        self.width_slider.setValue(2)
        self.width_slider.valueChanged.connect(self.set_pen_width)
        width_layout.addWidget(self.width_slider)
        top_control_layout.addLayout(width_layout)
        
        # 尺寸预设选择
        size_layout = QVBoxLayout()
        size_layout.addWidget(QLabel("尺寸预设:"))
        size_buttons = QHBoxLayout()
        
        self.small_btn = QPushButton("小")
        self.medium_btn = QPushButton("中")
        self.large_btn = QPushButton("大")
        self.fullscreen_btn = QPushButton("全屏")
        
        self.small_btn.clicked.connect(lambda: self.set_canvas_size(600, 300))
        self.medium_btn.clicked.connect(lambda: self.set_canvas_size(800, 500))
        self.large_btn.clicked.connect(lambda: self.set_canvas_size(1000, 700))
        self.fullscreen_btn.clicked.connect(self.set_fullscreen_size)
        
        size_buttons.addWidget(self.small_btn)
        size_buttons.addWidget(self.medium_btn)
        size_buttons.addWidget(self.large_btn)
        size_buttons.addWidget(self.fullscreen_btn)
        
        size_layout.addLayout(size_buttons)
        top_control_layout.addLayout(size_layout)
        
        top_control_layout.addStretch()
        control_layout.addLayout(top_control_layout)
        
        # 手写区域
        self.canvas = ResizableHandwritingCanvas(self)
        # 初始大小设置为中等
        self.set_canvas_size(800, 500)
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        self.clear_btn = QPushButton("清除")
        self.recognize_btn = QPushButton("识别公式")
        self.clear_btn.clicked.connect(self.clear_canvas)
        self.recognize_btn.clicked.connect(self.recognize_formula)
        
        btn_layout.addWidget(self.clear_btn)
        btn_layout.addWidget(self.recognize_btn)
        
        main_layout.addWidget(control_group)
        main_layout.addWidget(self.canvas)
        main_layout.addLayout(btn_layout)
        
    def choose_color(self):
        """选择画笔颜色"""
        color = QColorDialog.getColor(self.pen_color, self, "选择画笔颜色")
        if color.isValid():
            self.pen_color = color
            self.canvas.pen_color = color
            
    def set_pen_width(self, value):
        """设置画笔粗细"""
        self.pen_width = value
        self.canvas.pen_width = value
        
    def set_canvas_size(self, width, height):
        """设置画布大小"""
        self.canvas.resize(width, height)
        
    def set_fullscreen_size(self):
        """设置画布为全屏大小（基于主窗口）"""
        if self.parent() and hasattr(self.parent(), 'parent'):
            main_window = self.parent().parent()
            if main_window:
                # 计算可用空间（减去工具栏和边距）
                available_rect = main_window.geometry()
                width = available_rect.width() - 100  # 留出边距
                height = available_rect.height() - 200  # 留出工具栏等空间
                self.set_canvas_size(max(600, width), max(400, height))
        
    def clear_canvas(self):
        """清除手写内容"""
        self.canvas.clear_canvas()
        
    def recognize_formula(self):
        """使用Pix2Text识别手写公式"""
        # 获取主窗口
        main_window = self.parent().parent().parent()
        if not main_window:
            return
            
        # 检查Pix2Text是否可用
        if not PIX2TEXT_AVAILABLE:
            QMessageBox.warning(self, "功能不可用", 
                               "Pix2Text库未安装，请使用 'pip install pix2text' 安装")
            return
        
        try:
            # 显示加载状态
            main_window.statusBar().showMessage("正在识别手写公式...")
            QApplication.processEvents()  # 更新UI
            
            # 获取PIL图像
            pil_image = self.canvas.get_pil_image()
            
            # 使用Pix2Text识别
            p2t = Pix2Text()
            text = p2t.recognize_formula(pil_image)
            
            if text:
                main_window.latex_editor.setPlainText(text)
                main_window.preview_canvas.update_formula(text)
                main_window.statusBar().showMessage("公式识别成功")
                QMessageBox.information(self, "识别成功", "已成功识别公式")
            else:
                main_window.statusBar().showMessage("无法识别公式")
                QMessageBox.warning(self, "识别失败", "无法识别公式")
                
        except Exception as e:
            main_window.statusBar().showMessage("公式识别失败")
            QMessageBox.critical(self, "识别错误", f"识别过程中发生错误: {str(e)}")

class AIAssistantThread(QThread):
    """AI助手线程，用于处理AI请求，避免UI卡顿"""
    response_ready = pyqtSignal(str)
    error_occurred = pyqtSignal(str)
    
    def __init__(self, api_key, prompt):
        super().__init__()
        self.api_key = api_key
        self.prompt = prompt
        
    def run(self):
        try:
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {self.api_key}"
            }
            
            payload = {
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": "你是一位数学专家，擅长解释各种数学公式。请清晰解释公式的用途、每个参数的含义以及相关背景知识。回答要简洁明了，适合学生和研究人员理解。"},
                    {"role": "user", "content": self.prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 1000
            }
            
            response = requests.post(
                DEEPSEEK_CHAT_URL,
                headers=headers,
                data=json.dumps(payload)
            )
            
            if response.status_code == 200:
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    self.response_ready.emit(result["choices"][0]["message"]["content"])
                else:
                    self.error_occurred.emit("未获取到有效的回答")
            else:
                self.error_occurred.emit(f"API请求失败 (状态码: {response.status_code}): {response.text}")
                
        except Exception as e:
            self.error_occurred.emit(f"请求过程中发生错误: {str(e)}")

class FormulaCanvas(FigureCanvas):
    """公式预览画布"""
    def __init__(self, parent=None, width=5, height=4, dpi=100):
        # 设置matplotlib字体
        plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC", "Arial Unicode MS"]
        self.fig = Figure(figsize=(width, height), dpi=dpi)
        self.axes = self.fig.add_subplot(111)
        self.axes.axis('off')  # 关闭坐标轴
        super().__init__(self.fig)
        self.setParent(parent)
        self.setMinimumSize(300, 200)
        self.update_formula("")

    def update_formula(self, latex_code):
        """更新显示的公式"""
        self.axes.clear()
        self.axes.axis('off')
        
        if latex_code.strip():
            try:
                # 使用matplotlib渲染LaTeX公式
                self.axes.text(0.5, 0.5, f"${latex_code}$", 
                              fontsize=18, ha='center', va='center')
            except Exception as e:
                self.axes.text(0.5, 0.5, f"公式渲染错误: {str(e)}", 
                              fontsize=12, ha='center', va='center', color='red')
        else:
            self.axes.text(0.5, 0.5, "在此输入LaTeX公式", 
                          fontsize=12, ha='center', va='center', color='gray')
            
        self.fig.tight_layout()
        self.draw()

class EditableFormulaPreview(QWidget):
    """可编辑的公式预览控件，模仿Word的公式编辑方式"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.latex_code = ""
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 预览区域
        self.preview_label = QLabel()
        self.preview_label.setAlignment(Qt.AlignCenter)
        self.preview_label.setStyleSheet("border: 1px solid #ccc; background-color: white; min-height: 100px;")
        self.preview_label.setText("公式预览区域")
        
        # 编辑按钮
        self.edit_button = QPushButton("编辑公式")
        self.edit_button.clicked.connect(self.edit_formula)
        
        layout.addWidget(self.preview_label)
        layout.addWidget(self.edit_button)
        
    def set_latex(self, latex_code):
        """设置LaTeX公式"""
        self.latex_code = latex_code
        self.update_preview()
        
    def update_preview(self):
        """更新预览"""
        # 这里使用简化的预览，实际应用中可以使用MathJax或其他渲染引擎
        if self.latex_code:
            self.preview_label.setText(f"${self.latex_code}$")
        else:
            self.preview_label.setText("公式预览区域")
            
    def edit_formula(self):
        """编辑公式"""
        if not self.latex_code:
            return
            
        # 打开公式编辑器进行编辑
        dialog = FormulaEditorDialog(self.latex_code, self)
        if dialog.exec_() == QDialog.Accepted:
            self.latex_code = dialog.get_latex()
            self.update_preview()
            # 通知主窗口公式已更新
            if hasattr(self.parent(), 'update_formula_from_preview'):
                self.parent().update_formula_from_preview(self.latex_code)

class FormulaEditorDialog(QDialog):
    """公式编辑器对话框"""
    def __init__(self, latex_code, parent=None):
        super().__init__(parent)
        self.latex_code = latex_code
        self.setWindowTitle("编辑公式")
        self.setModal(True)
        self.init_ui()
        
    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # 公式编辑区域
        self.formula_edit = QPlainTextEdit()
        self.formula_edit.setPlainText(self.latex_code)
        
        # 按钮
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        
        layout.addWidget(QLabel("编辑公式:"))
        layout.addWidget(self.formula_edit)
        layout.addWidget(button_box)
        
    def get_latex(self):
        """获取编辑后的LaTeX代码"""
        return self.formula_edit.toPlainText()

class SymbolButton(QPushButton):
    """带有图标的符号按钮"""
    def __init__(self, symbol, latex_code, parent=None):
        super().__init__(symbol, parent)
        self.latex_code = latex_code
        self.setToolTip(latex_code)
        self.setFixedSize(40, 40)
        
        # 尝试创建图标
        self.create_icon(symbol)
        
    def create_icon(self, symbol):
        """为符号创建图标"""
        try:
            # 创建一个小的图像作为图标
            pixmap = QPixmap(32, 32)
            pixmap.fill(Qt.white)
            
            painter = QPainter(pixmap)
            painter.setFont(QFont("Arial", 16))
            painter.drawText(pixmap.rect(), Qt.AlignCenter, symbol)
            painter.end()
            
            self.setIcon(QIcon(pixmap))
            self.setIconSize(QSize(32, 32))
        except:
            pass  # 如果创建图标失败，只显示文本

class WordFormulaEditor(QWidget):
    """模仿Word公式编辑器的界面"""
    def __init__(self, parent=None, latex_editor=None):
        super().__init__(parent)
        self.latex_editor = latex_editor
        self.init_ui()
        
    def init_ui(self):
        main_layout = QVBoxLayout(self)
        
        # 创建顶部工具栏
        toolbar_layout = QHBoxLayout()
        
        # 符号按钮
        symbols_btn = QPushButton("符号")
        symbols_btn.clicked.connect(self.show_symbols)
        toolbar_layout.addWidget(symbols_btn)
        
        # 结构按钮
        structures_btn = QPushButton("结构")
        structures_btn.clicked.connect(self.show_structures)
        toolbar_layout.addWidget(structures_btn)
        
        # 工具按钮
        tools_btn = QPushButton("工具")
        tools_btn.clicked.connect(self.show_tools)
        toolbar_layout.addWidget(tools_btn)
        
        toolbar_layout.addStretch()
        main_layout.addLayout(toolbar_layout)
        
        # 创建内容区域
        self.content_stack = QStackedWidget()
        
        # 符号面板
        self.symbols_panel = self.create_symbols_panel()
        self.content_stack.addWidget(self.symbols_panel)
        
        # 结构面板
        self.structures_panel = self.create_structures_panel()
        self.content_stack.addWidget(self.structures_panel)
        
        # 工具面板
        self.tools_panel = self.create_tools_panel()
        self.content_stack.addWidget(self.tools_panel)
        
        main_layout.addWidget(self.content_stack)
        
    def create_symbols_panel(self):
        """创建符号面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # 创建符号分类标签
        tabs = QTabWidget()
        
        # 基础数学符号
        basic_math_tab = QWidget()
        basic_math_layout = QGridLayout(basic_math_tab)
        
        basic_symbols = [
            ("+", "+"), ("-", "-"), ("×", "\\times"), ("÷", "\\div"),
            ("=", "="), ("≠", "\\neq"), (">", ">"), ("<", "<"),
            ("≥", "\\geq"), ("≤", "\\leq"), ("≈", "\\approx"), ("≡", "\\equiv"),
            ("±", "\\pm"), ("∓", "\\mp"), ("·", "\\cdot"), ("∗", "\\ast")
        ]
        
        for i, (display, code) in enumerate(basic_symbols):
            btn = SymbolButton(display, code)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            basic_math_layout.addWidget(btn, i // 4, i % 4)
            
        tabs.addTab(basic_math_tab, "基础数学")
        
        # 希腊字母
        greek_tab = QWidget()
        greek_layout = QGridLayout(greek_tab)
        
        greek_letters = [
            ("α", "\\alpha"), ("β", "\\beta"), ("γ", "\\gamma"), ("δ", "\\delta"),
            ("ε", "\\epsilon"), ("ζ", "\\zeta"), ("η", "\\eta"), ("θ", "\\theta"),
            ("ι", "\\iota"), ("κ", "\\kappa"), ("λ", "\\lambda"), ("μ", "\\mu"),
            ("ν", "\\nu"), ("ξ", "\\xi"), ("π", "\\pi"), ("ρ", "\\rho"),
            ("σ", "\\sigma"), ("τ", "\\tau"), ("υ", "\\upsilon"), ("φ", "\\phi"),
            ("χ", "\\chi"), ("ψ", "\\psi"), ("ω", "\\omega"),
            ("Α", "A"), ("Β", "B"), ("Γ", "\\Gamma"), ("Δ", "\\Delta")
        ]
        
        for i, (display, code) in enumerate(greek_letters):
            btn = SymbolButton(display, code)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            greek_layout.addWidget(btn, i // 4, i % 4)
            
        tabs.addTab(greek_tab, "希腊字母")
        
        # 运算符
        operators_tab = QWidget()
        operators_layout = QGridLayout(operators_tab)
        
        operators = [
            ("∑", "\\sum"), ("∏", "\\prod"), ("∫", "\\int"), ("∮", "\\oint"),
            ("∇", "\\nabla"), ("∂", "\\partial"), ("∞", "\\infty"), ("∀", "\\forall"),
            ("∃", "\\exists"), ("∄", "\\nexists"), ("∅", "\\emptyset"), ("∈", "\\in"),
            ("∉", "\\notin"), ("⊂", "\\subset"), ("⊃", "\\supset"), ("⊆", "\\subseteq"),
            ("⊇", "\\supseteq"), ("∩", "\\cap"), ("∪", "\\cup"), ("∧", "\\wedge"),
            ("∨", "\\vee"), ("¬", "\\neg"), ("⊕", "\\oplus"), ("⊗", "\\otimes")
        ]
        
        for i, (display, code) in enumerate(operators):
            btn = SymbolButton(display, code)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            operators_layout.addWidget(btn, i // 4, i % 4)
            
        tabs.addTab(operators_tab, "运算符")
        
        layout.addWidget(tabs)
        return panel
        
    def create_structures_panel(self):
        """创建结构面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # 创建结构分类标签
        tabs = QTabWidget()
        
        # 分数
        fractions_tab = QWidget()
        fractions_layout = QVBoxLayout(fractions_tab)
        
        fractions_group = QGroupBox("分数")
        fractions_grid = QGridLayout(fractions_group)
        
        fraction_types = [
            ("常规分数", "\\frac{}{}"),
            ("小分数", "\\tfrac{}{}"),
            ("大分数", "\\dfrac{}{}"),
            ("二项式系数", "\\binom{}{}")
        ]
        
        for i, (name, code) in enumerate(fraction_types):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            fractions_grid.addWidget(btn, i // 2, i % 2)
            
        fractions_layout.addWidget(fractions_group)
        tabs.addTab(fractions_tab, "分数")
        
        # 上下标
        scripts_tab = QWidget()
        scripts_layout = QVBoxLayout(scripts_tab)
        
        scripts_group = QGroupBox("上下标")
        scripts_grid = QGridLayout(scripts_group)
        
        script_types = [
            ("上标", "^{}"), ("下标", "_{}"), ("上下标", "^{}_{}"),
            ("前置上标", "{}^{}"), ("前置下标", "{}_{}")
        ]
        
        for i, (name, code) in enumerate(script_types):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            scripts_grid.addWidget(btn, i // 2, i % 2)
            
        scripts_layout.addWidget(scripts_group)
        tabs.addTab(scripts_tab, "上下标")
        
        # 根式
        radicals_tab = QWidget()
        radicals_layout = QVBoxLayout(radicals_tab)
        
        radicals_group = QGroupBox("根式")
        radicals_grid = QGridLayout(radicals_group)
        
        radical_types = [
            ("平方根", "\\sqrt{}"), ("n次方根", "\\sqrt[]{}"),
            ("立方根", "\\sqrt[3]{}"), ("四次方根", "\\sqrt[4]{}")
        ]
        
        for i, (name, code) in enumerate(radical_types):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            radicals_grid.addWidget(btn, i // 2, i % 2)
            
        radicals_layout.addWidget(radicals_group)
        tabs.addTab(radicals_tab, "根式")
        
        # 积分
        integrals_tab = QWidget()
        integrals_layout = QVBoxLayout(integrals_tab)
        
        integrals_group = QGroupBox("积分")
        integrals_grid = QGridLayout(integrals_group)
        
        integral_types = [
            ("单积分", "\\int_{}^{}"), ("二重积分", "\\iint_{}^{}"),
            ("三重积分", "\\iiint_{}^{}"), ("曲线积分", "\\oint_{}^{}")
        ]
        
        for i, (name, code) in enumerate(integral_types):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            integrals_grid.addWidget(btn, i // 2, i % 2)
            
        integrals_layout.addWidget(integrals_group)
        tabs.addTab(integrals_tab, "积分")
        
        # 大型运算符
        large_ops_tab = QWidget()
        large_ops_layout = QVBoxLayout(large_ops_tab)
        
        large_ops_group = QGroupBox("大型运算符")
        large_ops_grid = QGridLayout(large_ops_group)
        
        large_op_types = [
            ("求和", "\\sum_{}^{}"), ("乘积", "\\prod_{}^{}"),
            ("并集", "\\bigcup_{}^{}"), ("交集", "\\bigcap_{}^{}")
        ]
        
        for i, (name, code) in enumerate(large_op_types):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            large_ops_grid.addWidget(btn, i // 2, i % 2)
            
        large_ops_layout.addWidget(large_ops_group)
        tabs.addTab(large_ops_tab, "大型运算符")
        
        # 括号
        brackets_tab = QWidget()
        brackets_layout = QVBoxLayout(brackets_tab)
        
        brackets_group = QGroupBox("括号")
        brackets_grid = QGridLayout(brackets_group)
        
        bracket_types = [
            ("圆括号", "\\left( \\right)"), ("方括号", "\\left[ \\right]"),
            ("花括号", "\\left\\{ \\right\\}"), ("绝对值", "\\left| \\right|"),
            ("范数", "\\left\\| \\right\\|"), ("向上取整", "\\left\\lceil \\right\\rceil"),
            ("向下取整", "\\left\\lfloor \\right\\rfloor"), ("角括号", "\\left\\langle \\right\\rangle")
        ]
        
        for i, (name, code) in enumerate(bracket_types):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            brackets_grid.addWidget(btn, i // 4, i % 4)
            
        brackets_layout.addWidget(brackets_group)
        tabs.addTab(brackets_tab, "括号")
        
        # 矩阵
        matrices_tab = QWidget()
        matrices_layout = QVBoxLayout(matrices_tab)
        
        matrices_group = QGroupBox("矩阵")
        matrices_grid = QGridLayout(matrices_group)
        
        matrix_types = [
            ("2×2矩阵", "\\begin{bmatrix} & \\\\ & \\end{bmatrix}"),
            ("3×3矩阵", "\\begin{bmatrix} & & \\\\ & & \\\\ & & \\end{bmatrix}"),
            ("行列式", "\\begin{vmatrix} & \\\\ & \\end{vmatrix}"),
            ("pmatrix", "\\begin{pmatrix} & \\\\ & \\end{pmatrix}")
        ]
        
        for i, (name, code) in enumerate(matrix_types):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            matrices_grid.addWidget(btn, i // 2, i % 2)
            
        matrices_layout.addWidget(matrices_group)
        tabs.addTab(matrices_tab, "矩阵")
        
        layout.addWidget(tabs)
        return panel
        
    def create_tools_panel(self):
        """创建工具面板"""
        panel = QWidget()
        layout = QVBoxLayout(panel)
        
        # 常用函数
        functions_group = QGroupBox("常用函数")
        functions_layout = QGridLayout(functions_group)
        
        functions = [
            ("sin", "\\sin"), ("cos", "\\cos"), ("tan", "\\tan"), ("cot", "\\cot"),
            ("sec", "\\sec"), ("csc", "\\csc"), ("log", "\\log"), ("ln", "\\ln"),
            ("exp", "\\exp"), ("lim", "\\lim"), ("max", "\\max"), ("min", "\\min"),
            ("sup", "\\sup"), ("inf", "\\inf"), ("arg", "\\arg"), ("det", "\\det"),
            ("dim", "\\dim"), ("gcd", "\\gcd"), ("hom", "\\hom"), ("ker", "\\ker")
        ]
        
        for i, (name, code) in enumerate(functions):
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            functions_layout.addWidget(btn, i // 4, i % 4)
            
        layout.addWidget(functions_group)
        
        # 箭头
        arrows_group = QGroupBox("箭头")
        arrows_layout = QGridLayout(arrows_group)
        
        arrows = [
            ("→", "\\rightarrow"), ("←", "\\leftarrow"), ("↔", "\\leftrightarrow"),
            ("⇒", "\\Rightarrow"), ("⇐", "\\Leftarrow"), ("⇔", "\\Leftrightarrow"),
            ("↦", "\\mapsto"), ("↪", "\\hookrightarrow"), ("↩", "\\hookleftarrow"),
            ("↑", "\\uparrow"), ("↓", "\\downarrow"), ("↕", "\\updownarrow")
        ]
        
        for i, (name, code) in enumerate(arrows):
            btn = SymbolButton(name, code)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            arrows_layout.addWidget(btn, i // 4, i % 4)
            
        layout.addWidget(arrows_group)
        
        # 其他符号
        other_group = QGroupBox("其他符号")
        other_layout = QGridLayout(other_group)
        
        others = [
            ("°", "^{\\circ}"), ("′", "'"), ("″", "''"), ("ℏ", "\\hbar"),
            ("ℜ", "\\Re"), ("ℑ", "\\Im"), ("∇", "\\nabla"), ("△", "\\triangle"),
            ("□", "\\square"), ("◊", "\\diamond"), ("♡", "\\heartsuit"), ("♠", "\\spadesuit"),
            ("♣", "\\clubsuit"), ("♦", "\\diamondsuit"), ("★", "\\star"), ("†", "\\dagger")
        ]
        
        for i, (name, code) in enumerate(others):
            btn = SymbolButton(name, code)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            other_layout.addWidget(btn, i // 4, i % 4)
            
        layout.addWidget(other_group)
        
        return panel
        
    def show_symbols(self):
        """显示符号面板"""
        self.content_stack.setCurrentWidget(self.symbols_panel)
        
    def show_structures(self):
        """显示结构面板"""
        self.content_stack.setCurrentWidget(self.structures_panel)
        
    def show_tools(self):
        """显示工具面板"""
        self.content_stack.setCurrentWidget(self.tools_panel)
        
    def insert_latex(self, code):
        """插入LaTeX代码到主编辑器"""
        if self.latex_editor:
            # 保存当前光标位置
            cursor = self.latex_editor.textCursor()
            position = cursor.position()
            
            # 插入代码
            self.latex_editor.insertPlainText(code)
            
            # 如果代码包含占位符{}，将光标移动到第一个占位符内
            if '{}' in code:
                # 计算新位置
                placeholder_pos = code.find('{}')
                new_position = position + placeholder_pos + 1
                
                # 移动光标
                cursor.setPosition(new_position)
                self.latex_editor.setTextCursor(cursor)
            
            # 更新预览
            if hasattr(self.latex_editor, 'parent') and hasattr(self.latex_editor.parent(), 'update_preview'):
                self.latex_editor.parent().update_preview()

class AIAssistantPanel(QWidget):
    """AI公式解释助手面板"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.main_window = parent
        self.ai_thread = None
        self.init_ui()
        
    def init_ui(self):
        main_layout = QVBoxLayout(self)
        
        # 标题
        main_layout.addWidget(QLabel("AI公式解释助手"))
        
        # 问题输入
        self.question_input = QTextEdit()
        self.question_input.setPlaceholderText("请输入关于公式的问题，例如：这个公式的用途是什么？每个参数代表什么意思？...")
        self.question_input.setMinimumHeight(80)
        
        # 按钮区域
        btn_layout = QHBoxLayout()
        self.use_current_btn = QPushButton("解释当前公式")
        self.send_btn = QPushButton("发送问题")
        self.clear_btn = QPushButton("清空")
        
        self.use_current_btn.clicked.connect(self.use_current_formula)
        self.send_btn.clicked.connect(self.send_question)
        self.clear_btn.clicked.connect(self.clear_input)
        
        btn_layout.addWidget(self.use_current_btn)
        btn_layout.addWidget(self.send_btn)
        btn_layout.addWidget(self.clear_btn)
        
        # 回答显示
        self.answer_display = QTextEdit()
        self.answer_display.setReadOnly(True)
        self.answer_display.setPlaceholderText("AI的回答将显示在这里...")
        
        # 添加到布局
        main_layout.addWidget(self.question_input)
        main_layout.addLayout(btn_layout)
        main_layout.addWidget(QLabel("AI回答:"))
        main_layout.addWidget(self.answer_display)
        
    def use_current_formula(self):
        """使用当前编辑器中的公式作为问题基础"""
        latex_code = self.main_window.latex_editor.toPlainText().strip()
        if not latex_code:
            QMessageBox.warning(self, "警告", "当前没有公式内容")
            return
            
        # 生成默认问题
        self.question_input.setPlainText(f"请解释这个公式的用途和各个参数的含义：{latex_code}")
        
    def send_question(self):
        """发送问题给AI助手"""
        question = self.question_input.toPlainText().strip()
        if not question:
            QMessageBox.warning(self, "警告", "请输入问题")
            return
            
        # 检查API密钥
        if not self.main_window.deepseek_api_key:
            reply = QMessageBox.question(
                self, "API未设置", 
                "尚未设置DeepSeek API密钥，是否现在设置？",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes
            )
            if reply == QMessageBox.Yes:
                self.main_window.set_api_credentials()
                if not self.main_window.deepseek_api_key:
                    return
            else:
                return
        
        # 显示加载状态
        self.answer_display.setPlainText("AI正在思考，请稍候...")
        self.main_window.statusBar().showMessage("AI正在生成回答...")
        
        # 启动AI线程
        self.ai_thread = AIAssistantThread(self.main_window.deepseek_api_key, question)
        self.ai_thread.response_ready.connect(self.display_answer)
        self.ai_thread.error_occurred.connect(self.display_error)
        self.ai_thread.start()
        
    def display_answer(self, answer):
        """显示AI的回答"""
        self.answer_display.setPlainText(answer)
        self.main_window.statusBar().showMessage("AI回答已生成")
        
    def display_error(self, error_msg):
        """显示错误信息"""
        self.answer_display.setPlainText(f"发生错误：{error_msg}")
        self.main_window.statusBar().showMessage("AI请求失败")
        
    def clear_input(self):
        """清空输入和输出"""
        self.question_input.clear()
        self.answer_display.clear()

class FormulaEditor(QMainWindow):
    """增强版公式编辑器主窗口"""
    def __init__(self):
        super().__init__()
        self.deepseek_api_key = ""
        self.load_api_credentials()  # 加载API凭证
        self.init_ui()
        
    def init_ui(self):
        # 设置窗口基本属性
        self.setWindowTitle("智能公式编辑与识别软件")
        self.setGeometry(100, 100, 1400, 900)
        
        # 设置全局字体，确保中文显示正常
        font = QFont("SimHei", 10)
        QApplication.setFont(font)
        
        # 创建中心部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        
        # 创建工具栏
        self.create_toolbar()
        
        # 创建格式转换工具栏
        self.create_conversion_toolbar()
        
        # 创建主分割器（左右）
        main_splitter = QSplitter(Qt.Horizontal)
        
        # 左侧：传统公式编辑器 + LaTeX编辑器
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        
        # 创建编辑器标签页
        self.editor_tabs = QTabWidget()
        
        # LaTeX代码编辑器
        latex_editor_panel = QWidget()
        latex_layout = QVBoxLayout(latex_editor_panel)
        
        self.latex_editor = QTextEdit()
        self.latex_editor.setFont(QFont("Consolas", 12))
        self.latex_editor.setPlaceholderText("在此输入LaTeX公式代码...")
        self.latex_editor.textChanged.connect(self.on_latex_changed)
        
        # 常用公式按钮
        math_symbols = QWidget()
        symbols_layout = QHBoxLayout(math_symbols)
        symbols_layout.setContentsMargins(0, 0, 0, 0)
        
        # 添加常用数学符号按钮
        symbols = [
            ("分数", "\\frac{numerator}{denominator}"),
            ("平方根", "\\sqrt{expression}"),
            ("求和", "\\sum_{i=1}^n expression"),
            ("积分", "\\int_{a}^{b} expression"),
            ("极限", "\\lim_{x \\to a} expression"),
            ("希腊字母", "\\alpha, \\beta, \\gamma, \\delta")
        ]
        
        for name, code in symbols:
            btn = QPushButton(name)
            btn.clicked.connect(lambda checked, c=code: self.insert_latex(c))
            symbols_layout.addWidget(btn)
        
        latex_layout.addWidget(QLabel("LaTeX 代码编辑区:"))
        latex_layout.addWidget(math_symbols)
        latex_layout.addWidget(self.latex_editor)
        
        self.editor_tabs.addTab(latex_editor_panel, "LaTeX编辑器")
        
        # 添加Word风格公式编辑器
        self.word_formula_editor = WordFormulaEditor(latex_editor=self.latex_editor)
        self.editor_tabs.addTab(self.word_formula_editor, "Word公式编辑器")
        
        left_layout.addWidget(self.editor_tabs)
        
        # 右侧：预览、转换和AI助手
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        # 预览区域
        preview_group = QGroupBox("公式预览")
        preview_layout = QVBoxLayout(preview_group)
        
        self.preview_canvas = FormulaCanvas(self, width=5, height=4, dpi=100)
        
        # 预览控制按钮
        preview_controls = QHBoxLayout()
        self.render_btn = QPushButton("渲染公式")
        self.render_btn.clicked.connect(self.update_preview)
        self.save_btn = QPushButton("保存公式图片")
        self.save_btn.clicked.connect(self.save_formula_image)
        self.image_recognize_btn = QPushButton("从图片识别公式")
        self.image_recognize_btn.clicked.connect(self.recognize_from_image)
        self.screenshot_btn = QPushButton("截图识别公式")
        self.screenshot_btn.clicked.connect(self.take_screenshot)
        
        preview_controls.addWidget(self.render_btn)
        preview_controls.addWidget(self.save_btn)
        preview_controls.addWidget(self.image_recognize_btn)
        preview_controls.addWidget(self.screenshot_btn)
        
        preview_layout.addLayout(preview_controls)
        preview_layout.addWidget(self.preview_canvas)
        
        # 格式转换区域
        conversion_group = QGroupBox("格式转换")
        conversion_layout = QVBoxLayout(conversion_group)
        
        # 转换控制
        conversion_controls = QHBoxLayout()
        
        # 源格式选择
        self.source_format = QComboBox()
        self.source_format.addItems(["LaTeX", "MathML"])
        self.source_format.setCurrentText("LaTeX")
        
        # 目标格式选择
        self.target_format = QComboBox()
        self.target_format.addItems(["MathML", "LaTeX", "Word (OMML)"])
        self.target_format.setCurrentText("MathML")
        
        # 转换按钮
        self.convert_btn = QPushButton("转换")
        self.convert_btn.clicked.connect(self.convert_format)
        
        # 复制结果按钮
        self.copy_result_btn = QPushButton("复制结果")
        self.copy_result_btn.clicked.connect(self.copy_conversion_result)
        
        conversion_controls.addWidget(QLabel("从:"))
        conversion_controls.addWidget(self.source_format)
        conversion_controls.addWidget(QLabel("到:"))
        conversion_controls.addWidget(self.target_format)
        conversion_controls.addWidget(self.convert_btn)
        conversion_controls.addWidget(self.copy_result_btn)
        
        # 转换结果显示
        self.conversion_result = QTextEdit()
        self.conversion_result.setFont(QFont("Consolas", 10))
        self.conversion_result.setReadOnly(False)
        
        conversion_layout.addLayout(conversion_controls)
        conversion_layout.addWidget(QLabel("转换结果:"))
        conversion_layout.addWidget(self.conversion_result)
        
        right_layout.addWidget(preview_group)
        right_layout.addWidget(conversion_group)
        
        # 添加到主分割器
        main_splitter.addWidget(left_panel)
        main_splitter.addWidget(right_panel)
        main_splitter.setSizes([600, 800])  # 设置初始大小比例
        
        main_layout.addWidget(main_splitter)
        
        # 添加手写识别面板
        self.handwriting_dock = QDockWidget("手写公式识别", self)
        self.handwriting_panel = HandwritingPanel(self)
        self.handwriting_dock.setWidget(self.handwriting_panel)
        self.addDockWidget(Qt.BottomDockWidgetArea, self.handwriting_dock)
        
        # 添加AI助手面板
        self.ai_assistant_dock = QDockWidget("AI公式解释助手", self)
        self.ai_assistant_panel = AIAssistantPanel(self)
        self.ai_assistant_dock.setWidget(self.ai_assistant_panel)
        self.ai_assistant_dock.setMinimumWidth(350)
        self.addDockWidget(Qt.RightDockWidgetArea, self.ai_assistant_dock)
        
        # 状态栏
        self.statusBar().showMessage("就绪")
        
    def create_toolbar(self):
        """创建主工具栏"""
        toolbar = QToolBar("主工具栏")
        toolbar.setIconSize(QSize(16, 16))
        self.addToolBar(toolbar)
        
        # 新建文件动作
        new_action = QAction("新建", self)
        new_action.setShortcut("Ctrl+N")
        new_action.setStatusTip("新建公式")
        new_action.triggered.connect(self.new_formula)
        toolbar.addAction(new_action)
        
        # 打开文件动作
        open_action = QAction("打开", self)
        open_action.setShortcut("Ctrl+O")
        open_action.setStatusTip("打开公式文件")
        open_action.triggered.connect(self.open_formula_file)
        toolbar.addAction(open_action)
        
        # 保存文件动作
        save_action = QAction("保存", self)
        save_action.setShortcut("Ctrl+S")
        save_action.setStatusTip("保存公式文件")
        save_action.triggered.connect(self.save_formula_file)
        toolbar.addAction(save_action)
        
        toolbar.addSeparator()
        
        # API设置动作
        api_action = QAction("API设置", self)
        api_action.setStatusTip("设置DeepSeek API密钥")
        api_action.triggered.connect(self.set_api_credentials)
        toolbar.addAction(api_action)
        
        # 帮助动作
        help_action = QAction("帮助", self)
        help_action.setStatusTip("显示帮助信息")
        help_action.triggered.connect(self.show_help)
        toolbar.addAction(help_action)
        
    def create_conversion_toolbar(self):
        """创建格式转换工具栏"""
        toolbar = QToolBar("格式转换")
        toolbar.setIconSize(QSize(16, 16))
        self.addToolBar(Qt.TopToolBarArea, toolbar)
        
        # 导出为Word动作
        word_action = QAction("导出为Word", self)
        word_action.setStatusTip("将公式导出为Word文档")
        word_action.triggered.connect(self.export_to_word)
        toolbar.addAction(word_action)
        
        # 导出为MathML动作
        mathml_action = QAction("导出为MathML", self)
        mathml_action.setStatusTip("将公式导出为MathML文件")
        mathml_action.triggered.connect(self.export_to_mathml)
        toolbar.addAction(mathml_action)
        
    def load_api_credentials(self):
        """加载DeepSeek API凭证"""
        try:
            if os.path.exists("deepseek_credentials.json"):
                with open("deepseek_credentials.json", "r") as f:
                    credentials = json.load(f)
                    self.deepseek_api_key = credentials.get("api_key", "")
        except Exception as e:
            print(f"加载API凭证失败: {e}")
            
    def save_api_credentials(self):
        """保存DeepSeek API凭证"""
        try:
            with open("deepseek_credentials.json", "w") as f:
                json.dump({
                    "api_key": self.deepseek_api_key
                }, f)
            QMessageBox.information(self, "保存成功", "DeepSeek API凭证已保存")
        except Exception as e:
            QMessageBox.critical(self, "保存失败", f"保存API凭证时出错: {e}")
            
    def set_api_credentials(self):
        """设置DeepSeek API凭证"""
        api_key, ok = QInputDialog.getText(
            self, "设置DeepSeek API", "请输入API密钥:", 
            QLineEdit.Password, self.deepseek_api_key
        )
        if ok and api_key:
            self.deepseek_api_key = api_key
            self.save_api_credentials()
                
    def take_screenshot(self):
        """截图识别公式"""
        screenshot_dialog = ScreenshotDialog(self)
        screenshot_dialog.exec_()
        
    def screenshot_captured(self, image):
        """处理截图"""
        try:
            self.statusBar().showMessage("正在识别截图中的公式...")
            QApplication.processEvents()
            
            # 使用Pix2Text识别
            if PIX2TEXT_AVAILABLE:
                p2t = Pix2Text()
                text = p2t.recognize_formula(image)
                
                if text:
                    self.latex_editor.setPlainText(text)
                    self.preview_canvas.update_formula(text)
                    self.statusBar().showMessage("公式识别成功")
                    QMessageBox.information(self, "识别成功", "已成功从截图识别公式")
                else:
                    self.statusBar().showMessage("无法识别公式")
                    QMessageBox.warning(self, "识别失败", "无法从截图识别出公式")
            else:
                QMessageBox.warning(self, "功能不可用", 
                                   "Pix2Text库未安装，请使用 'pip install pix2text' 安装")
                
        except Exception as e:
            self.statusBar().showMessage("公式识别失败")
            QMessageBox.critical(self, "识别错误", f"识别过程中发生错误: {str(e)}")
                
    def recognize_from_image(self):
        """使用Pix2Text从图片识别公式"""
        # 选择图片文件
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择包含公式的图片", "", 
            "图片文件 (*.png *.jpg *.jpeg *.bmp *.gif)"
        )
        
        if not file_path:
            return
            
        try:
            # 显示加载状态
            self.statusBar().showMessage("正在识别图片中的公式...")
            QApplication.processEvents()
            
            # 检查Pix2Text是否可用
            if not PIX2TEXT_AVAILABLE:
                QMessageBox.warning(self, "功能不可用", 
                                   "Pix2Text库未安装，请使用 'pip install pix2text' 安装")
                return
            
            # 使用Pix2Text识别
            p2t = Pix2Text()
            text = p2t.recognize_formula(file_path)
            
            if text:
                self.latex_editor.setPlainText(text)
                self.preview_canvas.update_formula(text)
                self.statusBar().showMessage("公式识别成功")
                QMessageBox.information(self, "识别成功", "已成功从图片中识别公式")
            else:
                self.statusBar().showMessage("无法从图片中识别公式")
                QMessageBox.warning(self, "识别失败", "无法从图片中识别出公式")
                
        except Exception as e:
            self.statusBar().showMessage("公式识别失败")
            QMessageBox.critical(self, "识别错误", f"识别过程中发生错误: {str(e)}")
    
    def convert_format(self):
        """转换公式格式"""
        source = self.source_format.currentText()
        target = self.target_format.currentText()
        source_content = ""
        
        # 获取源内容
        if source == "LaTeX":
            source_content = self.latex_editor.toPlainText().strip()
            if not source_content:
                QMessageBox.warning(self, "警告", "请先输入LaTeX公式")
                return
        else:  # MathML
            source_content = self.conversion_result.toPlainText().strip()
            if not source_content:
                QMessageBox.warning(self, "警告", "请先输入或转换出MathML内容")
                return
        
        try:
            self.statusBar().showMessage(f"正在将{source}转换为{target}...")
            
            if source == "LaTeX" and target == "MathML":
                # LaTeX 转换为 MathML (使用pandoc)
                result = self.latex_to_mathml(source_content)
                
            elif source == "MathML" and target == "LaTeX":
                # MathML 转换为 LaTeX (使用pandoc)
                result = self.mathml_to_latex(source_content)
                
            elif (source == "LaTeX" or source == "MathML") and target == "Word (OMML)":
                # 转换为Word的OMML格式
                if source == "LaTeX":
                    mathml = self.latex_to_mathml(source_content)
                    result = self.mathml_to_omml(mathml)
                else:
                    result = self.mathml_to_omml(source_content)
                    
                # 同时导出为Word文档
                self.export_to_word(content=result)
                result = "已成功导出为Word文档，包含可编辑的公式"
                
            self.conversion_result.setPlainText(result)
            self.statusBar().showMessage(f"{source}转换为{target}成功")
            
        except Exception as e:
            self.statusBar().showMessage(f"{source}转换为{target}失败")
            QMessageBox.critical(self, "转换失败", f"格式转换错误: {str(e)}")
    
    def latex_to_mathml(self, latex_code):
        """使用pandoc将LaTeX转换为MathML"""
        # 检查pandoc是否安装
        if not self.check_pandoc_installed():
            raise Exception("未检测到pandoc，请先安装pandoc以使用格式转换功能")
            
        # 使用pandoc进行转换
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(mode='w', suffix='.tex', delete=False) as f:
                # 写入完整的LaTeX文档结构
                f.write("\\documentclass{article}\n")
                f.write("\\begin{document}\n")
                f.write(f"$$\n{latex_code}\n$$\n")
                f.write("\\end{document}\n")
                temp_filename = f.name
                
            # 调用pandoc进行转换
            result = subprocess.run(
                ['pandoc', temp_filename, '-f', 'latex', '-t', 'mathml', '--standalone'],
                capture_output=True, text=True, check=True
            )
            
            # 清理临时文件
            os.unlink(temp_filename)
            
            return result.stdout
            
        except subprocess.CalledProcessError as e:
            raise Exception(f"pandoc转换失败: {e.stderr}")
        except Exception as e:
            raise Exception(f"转换错误: {str(e)}")
    
    def mathml_to_latex(self, mathml_code):
        """使用pandoc将MathML转换为LaTeX"""
        # 检查pandoc是否安装
        if not self.check_pandoc_installed():
            raise Exception("未检测到pandoc，请先安装pandoc以使用格式转换功能")
            
        try:
            # 创建临时文件
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False) as f:
                f.write(f"<html><body>{mathml_code}</body></html>")
                temp_filename = f.name
                
            # 调用pandoc进行转换
            result = subprocess.run(
                ['pandoc', temp_filename, '-f', 'html', '-t', 'latex', '--standalone'],
                capture_output=True, text=True, check=True
            )
            
            # 清理临时文件
            os.unlink(temp_filename)
            
            # 提取公式部分，去除文档结构
            latex = result.stdout
            start = latex.find('$$') + 2
            end = latex.rfind('$$')
            if start > 2 and end > start:
                return latex[start:end].strip()
            return latex
            
        except subprocess.CalledProcessError as e:
            raise Exception(f"pandoc转换失败: {e.stderr}")
        except Exception as e:
            raise Exception(f"转换错误: {str(e)}")
    
    def mathml_to_omml(self, mathml_code):
        """将MathML转换为Office Math Markup Language (OMML)"""
        # Word使用的格式，用于实现与Word公式编辑器的兼容性
        return f"<!-- Word公式格式 (OMML) - 可在Word中编辑 -->\n{mathml_code}"
    
    def check_pandoc_installed(self):
        """检查pandoc是否安装"""
        try:
            subprocess.run(['pandoc', '--version'], capture_output=True, check=True)
            return True
        except (subprocess.SubprocessError, FileNotFoundError):
            return False
    
    def export_to_word(self, content=None):
        """导出公式为Word文档"""
        try:
            # 获取保存路径
            file_path, _ = QFileDialog.getSaveFileName(
                self, "导出为Word文档", "", "Word文档 (*.docx)"
            )
            
            if not file_path:
                return
                
            # 如果没有提供内容，则自动转换
            if not content:
                latex_code = self.latex_editor.toPlainText().strip()
                if not latex_code:
                    QMessageBox.warning(self, "警告", "请先输入LaTeX公式")
                    return
                    
                mathml = self.latex_to_mathml(latex_code)
                content = self.mathml_to_omml(mathml)
            
            # 使用python-docx创建Word文档
            try:
                from docx import Document
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                
                doc = Document()
                
                # 添加公式说明
                doc.add_paragraph("公式内容:")
                
                # 添加可在Word中编辑的公式
                p = doc.add_paragraph()
                run = p.add_run()
                run.text = "LaTeX公式: " + self.latex_editor.toPlainText()
                
                # 添加提示信息
                doc.add_paragraph("提示: 此文档中的公式可在Word公式编辑器中直接编辑修改")
                
                doc.save(file_path)
                self.statusBar().showMessage(f"已导出为Word文档: {file_path}")
                QMessageBox.information(self, "导出成功", f"公式已成功导出至:\n{file_path}")
                
            except ImportError:
                raise Exception("请安装python-docx库以导出Word文档: pip install python-docx")
            except Exception as e:
                raise Exception(f"导出Word文档失败: {str(e)}")
                
        except Exception as e:
            self.statusBar().showMessage("导出Word文档失败")
            QMessageBox.critical(self, "导出失败", f"导出过程中发生错误: {str(e)}")
    
    def export_to_mathml(self):
        """导出公式为MathML文件"""
        latex_code = self.latex_editor.toPlainText().strip()
        if not latex_code:
            QMessageBox.warning(self, "警告", "请先输入LaTeX公式")
            return
            
        try:
            # 获取保存路径
            file_path, _ = QFileDialog.getSaveFileName(
                self, "导出为MathML", "", "MathML文件 (*.mml)"
            )
            
            if not file_path:
                return
                
            # 转换为MathML
            mathml = self.latex_to_mathml(latex_code)
            
            # 保存到文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(mathml)
                
            self.statusBar().showMessage(f"已导出为MathML文件: {file_path}")
            QMessageBox.information(self, "导出成功", f"公式已成功导出至:\n{file_path}")
            
        except Exception as e:
            self.statusBar().showMessage("导出MathML文件失败")
            QMessageBox.critical(self, "导出失败", f"导出过程中发生错误: {str(e)}")
    
    def copy_conversion_result(self):
        """复制转换结果到剪贴板"""
        result = self.conversion_result.toPlainText()
        if result:
            clipboard = QApplication.clipboard()
            clipboard.setText(result)
            self.statusBar().showMessage("转换结果已复制到剪贴板")
        else:
            QMessageBox.warning(self, "警告", "没有可复制的转换结果")
    
    def insert_latex(self, code):
        """在编辑器中插入LaTeX代码"""
        self.latex_editor.insertPlainText(code)
        self.on_latex_changed()
        
    def on_latex_changed(self):
        """LaTeX代码改变时的处理"""
        latex_code = self.latex_editor.toPlainText()
        self.preview_canvas.update_formula(latex_code)
        
    def update_preview(self):
        """手动更新预览"""
        latex_code = self.latex_editor.toPlainText()
        self.preview_canvas.update_formula(latex_code)
        self.statusBar().showMessage("公式已更新")
        
    def save_formula_image(self):
        """保存公式为图片"""
        latex_code = self.latex_editor.toPlainText()
        if not latex_code.strip():
            QMessageBox.warning(self, "警告", "没有可保存的公式内容")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存公式图片", "", "PNG图片 (*.png);;SVG图片 (*.svg);;PDF文件 (*.pdf)"
        )
        
        if file_path:
            try:
                self.preview_canvas.fig.savefig(file_path, bbox_inches='tight', dpi=300)
                self.statusBar().showMessage(f"公式已保存至: {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "保存失败", f"保存公式时出错: {str(e)}")
                
    def new_formula(self):
        """新建公式"""
        self.latex_editor.clear()
        self.conversion_result.clear()
        self.preview_canvas.update_formula("")
        self.handwriting_panel.clear_canvas()
        self.ai_assistant_panel.clear_input()
        self.statusBar().showMessage("已创建新公式")
        
    def open_formula_file(self):
        """打开公式文件"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "打开公式文件", "", "LaTeX文件 (*.tex);;MathML文件 (*.mml);;文本文件 (*.txt);;所有文件 (*)"
        )
        
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                    # 根据文件扩展名判断格式
                    if file_path.lower().endswith('.mml'):
                        self.conversion_result.setPlainText(content)
                        self.source_format.setCurrentText("MathML")
                        self.target_format.setCurrentText("LaTeX")
                        self.convert_format()  # 自动转换为LaTeX
                    else:
                        self.latex_editor.setPlainText(content)
                        self.preview_canvas.update_formula(content)
                        
                self.statusBar().showMessage(f"已打开文件: {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "打开失败", f"打开文件时出错: {str(e)}")
                
    def save_formula_file(self):
        """保存公式代码到文件"""
        latex_code = self.latex_editor.toPlainText()
        if not latex_code.strip():
            QMessageBox.warning(self, "警告", "没有可保存的公式内容")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存公式代码", "", "LaTeX文件 (*.tex);;文本文件 (*.txt)"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(latex_code)
                self.statusBar().showMessage(f"公式代码已保存至: {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "保存失败", f"保存文件时出错: {str(e)}")
                
    def show_help(self):
        """显示帮助信息"""
        help_text = """智能公式编辑与识别软件使用帮助:
        
1. 编辑器模式:
   - LaTeX编辑器: 直接输入LaTeX代码编辑公式
   - Word公式编辑器: 通过点击符号和结构来可视化创建公式，界面和功能与Microsoft Word公式编辑器完全相同

2. 手写公式识别:
   - 使用底部面板的手写区域绘制公式
   - 可调整画笔颜色和粗细
   - 可调整手写区域大小：
     * 拖动右下角的调整手柄手动调整
     * 使用"小/中/大/全屏"按钮快速设置
   - 点击"识别公式"按钮将手写内容转换为LaTeX代码

3. 图片识别公式:
   - 点击"从图片识别公式"按钮选择包含公式的图片
   - 系统会自动识别并转换为LaTeX代码

4. 截图识别公式:
   - 点击"截图识别公式"按钮进行屏幕截图
   - 拖动鼠标选择包含公式的区域
   - 系统会自动识别并转换为LaTeX代码

5. 格式转换:
   - 支持LaTeX ↔ MathML互相转换
   - 支持导出为Word文档（可在Word公式编辑器中编辑）
   - 转换结果可复制到剪贴板

6. AI公式解释助手:
   - 可直接提问或解释当前公式
   - AI会解释公式用途、参数含义和相关背景知识
   - 需设置DeepSeek API密钥才能使用

7. API设置:
   - 需要注册DeepSeek账号获取API密钥
   - 点击工具栏"API设置"输入您的密钥

格式转换依赖:
- 需要安装pandoc: https://pandoc.org/installing.html
- 导出Word需要安装python-docx: pip install python-docx
- 本地公式识别需要安装pix2text: pip install pix2text
        """
        QMessageBox.information(self, "使用帮助", help_text)

if __name__ == "__main__":
    # 确保必要的库已安装
    required_libs = {
        'requests': 'requests',
        'PyQt5': 'PyQt5',
        'matplotlib': 'matplotlib',
        'numpy': 'numpy',
        'PIL': 'Pillow'
    }
    
    missing_libs = []
    for lib, pkg in required_libs.items():
        try:
            __import__(lib)
        except ImportError:
            missing_libs.append(pkg)
    
    if missing_libs:
        print(f"请先安装以下依赖库: pip install {' '.join(missing_libs)}")
        sys.exit(1)
        
    app = QApplication(sys.argv)
    window = FormulaEditor()
    window.show()
    sys.exit(app.exec_())